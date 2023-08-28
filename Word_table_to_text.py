from docx import Document

def extract_column_from_table(docx_path, table_index, column_index):
    document = Document(docx_path)
    
    if table_index < len(document.tables):
        table = document.tables[table_index]
        column_content = []
        
        for row in table.rows:
            if column_index < len(row.cells):
                cell = row.cells[column_index]
                column_content.append(cell.text.strip())  # Remove leading/trailing spaces and newline characters
        
        return column_content
    else:
        return None

def write_list_to_text_file(data_list, output_file):
    with open(output_file, 'w', encoding='utf-8') as file:
        for item in data_list:
            file.write("%s\n" % item)

if __name__ == "__main__":
    input_docx_path1 = "C:\\Users\\st_je\\OneDrive\\Documents\\Smartcat\\Tools\\BLEU Scores\\HBEA and Public Charge(en-es).docx"
    input_docx_path2 = "C:\\Users\\st_je\\OneDrive\\Documents\\Smartcat\\Tools\\BLEU Scores\\HBEA and Public Charge.docx"
    table_index1 = 0  # Index of the table in the first Word file
    table_index2 = 0  # Index of the table in the second Word file
    column_index1 = 2  # Index of the column in the first table (0-based)
    column_index2 = 2  # Index of the column in the second table (0-based)
    output_text_file = "C:\\Users\\st_je\\OneDrive\\Documents\\Smartcat\\Tools\\BLEU Scores\\output2.txt"
    
    column_content1 = extract_column_from_table(input_docx_path1, table_index1, column_index1)
    column_content2 = extract_column_from_table(input_docx_path2, table_index2, column_index2)
    
    if column_content1 and column_content2:
        combined_content = [f"{item1}\t{item2}" for item1, item2 in zip(column_content1, column_content2)]
        write_list_to_text_file(combined_content, output_text_file)
        print("Combined content extracted and written to", output_text_file)
    else:
        print("Table not found in one or both documents, or invalid indices.")

